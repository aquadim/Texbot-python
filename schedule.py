# schedule.py
# Модуль для работы с расписанием. Не используется в главном модуле

import requests
import os
import subprocess
import datetime
import database
import sys
import time
import hashlib
from docx import Document
from utils import *

# Статус функции downloadAndCheck
ALL_FOUND = 0
FOUND_TODAY = 1
FOUND_TOMORROW = 2
NOT_FOUND = 3

def downloadScheduleFile(name):
	try:
		return requests.get(name)
	except Exception as e:
		print2(e, 'red')
		return False

def getDateList(word_doc, now):
	"""Возвращает список дат, соответствующий таблицам расписаний.
	Если элемент списка False, то расписание не актуально"""
	output = []

	# Проверка параграфов текста на привязку к таблицам
	for item in word_doc.paragraphs:
		paragraph = item.text.lower()
		words = paragraph.split(" ")

		# Должно быть не менее 4 слов
		if len(words) < 4:
			continue

		# Поиск месяца
		month = -1
		for m in gen_month_num_to_str:
			if paragraph.find(gen_month_num_to_str[m]) != -1:
				month = m
				break
		if month == -1:
			# Не найдено совпадений для месяца
			output.append(False)
			continue

		# Поиск дня
		if words[4].isdigit():
			day = int(words[4])
		else:
			# Не обнаружен день
			output.append(False)
			continue

		# Вычисление года
		# Получаем дату и узнаём день недели этой даты
		# Если у этой даты и даты в документе день недели одинаковый, то мы нашли год!
		max_guess_attempts = 2
		for i in range(max_guess_attempts):
			guessed_date = datetime.datetime(year=now.year + i, month=month, day=day)
			guessed_weekday = guessed_date.isoweekday()

			if words[3] == gen_weekdays_num_to_str[guessed_weekday].lower():
				# Совпадает! Значит высока вероятность того, что год угадан правильно
				output.append(f'{now.year + i}-{dd(month)}-{dd(day)}')
				break

		output.append(False)

	return output

def parseTable(table, date):
	"""Парсит таблицу"""
	rows = table.rows
	height = len(rows)
	width = len(rows[0].cells)

	# Проверяем точное ли это расписание или нет
	is_uncertain = rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb != None
	if is_uncertain:
		print2(f"Таблица на дату {date} имеет неточные данные!", "red")
		return False

	# Перенос данных таблицы в двоичный массив
	data = []
	for y in range(height):
		data_row = []
		table_row = rows[y].cells
		for x in range(width):
			if x % 2 == 0:
				data_row.append(table_row[x].text.replace("\n", "").replace("\xa0", "").replace(".", ":"))
			else:
				data_row.append(table_row[x].text.replace("\n", "").replace("\xa0", ""))
		data.append(data_row)

	# Парсинг таблицы
	current_groups_row = [] # Какие группы в курсе сейчас присутствуют (в списке id групп)
	y = 0
	while y < len(data):
		if isGroupName(data[y][0]):
			current_groups_row.clear()
			for x in range(0, len(data[y]), 2):
				current_groups_row.append(database.cmdGetGidFromString(data[y][x]))
			y += 1
		else:
			for x in range(0, len(data[y]) - 1, 2):
				# Проверка правильности
				if len(data[y][x + 1]) < 2:
					# Названия пар не бывают такими короткими
					continue
				if len(data[y][x]) < 3:
					# Время не бывает настолько коротким
					continue

				# Узнаём id расписания для этой пары
				schedule_id = database.getScheduleId(current_groups_row[x // 2], date)
				if not schedule_id:
					# Этого расписания ещё не было создано. Создаём!
					schedule_id = database.addSchedule(current_groups_row[x // 2], date)
				else:
					# Это расписание уже существует. Очищаем все его пары если можно
					if database.getIfCanCleanSchedule(schedule_id):
						database.cleanSchedule(schedule_id)

				# Время пары
				pair_time = data[y][x].split(':')
				pair_time = dd(int(pair_time[0])) + ':' + dd(int(pair_time[1]))

				# Добавляем пару
				pair_id = database.addPair(schedule_id, pair_time, y, data[y][x+1])

				# К паре добавляем места пары
				# Место пары может быть в двух местах, в таких случаях места разделяются слэшем
				places_data = data[y+1][x+1].split('/')
				for index, place in enumerate(places_data):
					components = place.split(' ')
					if len(components) == 2:
						# Есть и кабинет и преподаватель
						database.addPairPlace(pair_id, database.getTeacherId(components[0]), components[1])
					else:
						# Есть только преподаватель
						database.addPairPlace(pair_id, database.getTeacherId(components[0]), None)
			y += 2
	print2(f'Таблица на дату {date} готова!', 'green')

def parseDocument(word_doc, date_index, date, now):
	"""Парсит документ"""
	parseTable(word_doc.tables[date_index], date)
	database.db.commit()
	end_time = time.time()

def downloadAndCheck(f, now, today, tomorrow, __dir__):
	"""Скачивает файл расписания и проверяет его на нужные нам даты"""
	# Скачиваем файл
	print(f"Скачиваем {f}")
	s = downloadScheduleFile(f)
	if not s:
		print2(f"Не удалось скачать f", "red")
		return (NOT_FOUND, word_doc)
	with open(__dir__ + "/tmp/schedule.doc", "wb") as f:
		f.write(s.content)

	# Конвертируем в docx
	subprocess.call(["lowriter", "--convert-to", "docx", __dir__+"/tmp/schedule.doc", "--outdir", __dir__+"/tmp/"])

	# Смотрим какие даты у этого документа есть
	word_doc = Document(__dir__+"/tmp/schedule.docx")
	table_dates = getDateList(word_doc, now)

	try:
		today_index = table_dates.index(today)
		has_today = True
	except ValueError:
		has_today = False

	try:
		tomorrow_index = table_dates.index(tomorrow)
		has_tomorrow = True
	except ValueError:
		has_tomorrow = False

	if has_today and has_tomorrow:
		print2("Обе даты найдены", "green")
		return (ALL_FOUND, word_doc, today_index, tomorrow_index)

	elif not has_today and has_tomorrow:
		print2("Завтра найдено", "green")
		return (FOUND_TOMORROW, word_doc, tomorrow_index)

	elif has_today and not has_tomorrow:
		print2("Сегодня найдено", "green")
		return (FOUND_TODAY, word_doc, today_index)

	else:
		print2("Ничего не найдено", "red")
		return (NOT_FOUND, word_doc)

def updateSchedule(__dir__):
	"""Загружает расписание, обновляет его в БД"""

	# Составляем даты которые нам нужны на сегодня
	now = datetime.datetime.now()
	date_today = now.strftime("%Y-%m-%d")
	date_tomorrow = (now + datetime.timedelta(days=1)).strftime("%Y-%m-%d")

	# Загрузка файла расписания из файла специально для Техбота
	response = downloadAndCheck("http://www.vpmt.ru/docs/rasp2.doc", now, date_today, date_tomorrow, __dir__)
	status = response[0]
	word_doc = response[1]

	if status == ALL_FOUND:
		parseDocument(word_doc, response[2], date_today, now)
		parseDocument(word_doc, response[3], date_tomorrow, now)
		exit()
	elif status == FOUND_TODAY:
		parseDocument(word_doc, response[2], date_today, now)
	elif status == FOUND_TOMORROW:
		parseDocument(word_doc, response[2], date_tomorrow, now)
	else:
		pass

	# Если программа ещё не вышла, то мы не нашли все даты. Попытка #2 с другим файлом
	old_status = status
	response = downloadAndCheck("http://www.vpmt.ru/docs/rasp.doc", now, date_today, date_tomorrow, __dir__)
	status = response[0]
	word_doc = response[1]
	
	if status == ALL_FOUND:
		if old_status == NOT_FOUND:
			# В прошлой попытке ничего не нашли, генерируем всё заново
			parseDocument(word_doc, response[2], date_today, now)
			parseDocument(word_doc, response[3], date_tomorrow, now)
		elif old_status == FOUND_TODAY:
			# В прошлой попытке уже нашли на сегодня, не делаем на сегодня
			parseDocument(word_doc, response[2], date_today, now)
		elif old_status == FOUND_TOMORROW:
			# В прошлой попытке уже нашли на сегодня, не делаем на сегодня
			parseDocument(word_doc, response[2], date_tomorrow, now)			
		exit()
		
	elif status == FOUND_TODAY:
		# Если в новой попытке мы снова нашли только сегодня, то не выполняем расчёты, заканчиваем программу
		if old_status == FOUND_TODAY:
			exit()
		parseDocument(word_doc, response[2], date_today, now)
	
	elif status == FOUND_TOMORROW:
		# Если в новой попытке мы снова нашли только завтра, то не выполняем расчёты, заканчиваем программу
		if old_status == FOUND_TOMORROW:
			exit()
		parseDocument(word_doc, response[2], date_tomorrow, now)

	database.makeSchedulesCleanable()
	database.db.close()

if __name__ == "__main__":
	__dir__ = os.path.dirname(__file__)
	updateSchedule(__dir__)
